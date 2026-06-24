from flask import Flask, render_template, request, redirect, session, send_file
from database import init_db, get_db
import psycopg2.extras
import pandas as pd
from docx import Document

app = Flask(__name__)
app.secret_key = "erp_pro_max"

init_db()

# ================= LOGIN =================
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        session["user"] = request.form["username"]
        return redirect("/dashboard")
    return render_template("login.html")


# ================= DASHBOARD =================
@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect("/")

    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) FROM products")
    total_items = cur.fetchone()[0] or 0

    cur.execute("SELECT COALESCE(SUM(quantity),0) FROM products")
    total_stock = cur.fetchone()[0] or 0

    cur.execute("SELECT COUNT(*) FROM products WHERE quantity < 5")
    low_stock = cur.fetchone()[0] or 0

    # FIX: không dùng total (tránh lỗi DB)
    cur.execute("""
        SELECT COALESCE(SUM(stock_log.quantity * products.price),0)
        FROM stock_log
        JOIN products ON products.id = stock_log.product_id
        WHERE stock_log.action = 'SELL'
    """)
    revenue = cur.fetchone()[0] or 0

    cur.execute("""
        SELECT category, COUNT(*)
        FROM products
        GROUP BY category
    """)

    rows = cur.fetchall()

    pie_labels = [r[0] for r in rows]
    pie_values = [r[1] for r in rows]

    conn.close()

    return render_template(
        "dashboard.html",
        total_items=total_items,
        total_stock=total_stock,
        low_stock=low_stock,
        today_income=revenue,
        pie_labels=pie_labels,
        pie_values=pie_values
    )


# ================= PRODUCTS =================
@app.route("/products", methods=["GET", "POST"])
def products():
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    if request.method == "POST":
        name = request.form["name"]
        category = request.form["category"]
        quantity = int(request.form["quantity"])
        price = int(request.form["price"])

        cur.execute("""
            INSERT INTO products(name, category, quantity, price)
            VALUES (%s,%s,%s,%s)
        """, (name, category, quantity, price))

    cur.execute("SELECT * FROM products ORDER BY id DESC")
    data = cur.fetchall()

    conn.commit()
    conn.close()

    return render_template("products.html", products=data)


# ================= SELL =================
@app.route("/sell/<int:id>", methods=["POST"])
def sell(id):
    conn = get_db()
    cur = conn.cursor()

    qty = int(request.form["quantity"])

    cur.execute("SELECT quantity, price FROM products WHERE id=%s", (id,))
    item = cur.fetchone()

    if item and item[0] >= qty:
        price = item[1]

        cur.execute("""
            UPDATE products
            SET quantity = quantity - %s
            WHERE id=%s
        """, (qty, id))

        cur.execute("""
            INSERT INTO stock_log(product_id, action, quantity)
            VALUES (%s,'SELL',%s)
        """, (id, qty))

    conn.commit()
    conn.close()
    return redirect("/products")


# ================= IMPORT =================
@app.route("/import/<int:id>", methods=["POST"])
def import_stock(id):
    conn = get_db()
    cur = conn.cursor()

    qty = int(request.form["quantity"])

    cur.execute("""
        UPDATE products
        SET quantity = quantity + %s
        WHERE id=%s
    """, (qty, id))

    cur.execute("""
        INSERT INTO stock_log(product_id, action, quantity)
        VALUES (%s,'IMPORT',%s)
    """, (id, qty))

    conn.commit()
    conn.close()
    return redirect("/products")


# ================= HISTORY (SAFE) =================
@app.route("/history")
def history():
    conn = get_db()
    cur = conn.cursor()

    cur.execute("""
        SELECT
            stock_log.id,
            products.name,
            stock_log.action,
            stock_log.quantity,
            stock_log.created_at
        FROM stock_log
        JOIN products ON products.id = stock_log.product_id
        ORDER BY stock_log.id DESC
    """)

    rows = cur.fetchall()
    conn.close()

    return render_template("history.html", logs=rows)


# ================= EXPORT EXCEL =================
@app.route("/export/excel")
def export_excel():
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT * FROM products")
    data = cur.fetchall()

    pd.DataFrame(data).to_excel("products.xlsx", index=False)

    return send_file("products.xlsx", as_attachment=True)


# ================= EXPORT WORD =================
@app.route("/export/word")
def export_word():
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT * FROM products")
    data = cur.fetchall()

    doc = Document()
    doc.add_heading("Báo cáo kho", 0)

    for r in data:
        doc.add_paragraph(f"{r['name']} | SL:{r['quantity']} | Giá:{r['price']}")

    doc.save("products.docx")

    return send_file("products.docx", as_attachment=True)


# ================= RUN =================
if __name__ == "__main__":
    app.run(debug=True)